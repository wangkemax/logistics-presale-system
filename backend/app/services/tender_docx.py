"""Tender Word document generator.

Produces a professional .docx proposal with:
- Cover page
- Table of contents
- 10 content chapters auto-filled from pipeline outputs
- Appendices
"""

import io
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

import structlog

logger = structlog.get_logger()


def _set_cell_shading(cell, color: str):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color,
        qn("w:val"): "clear",
    })
    shading.append(shading_elm)


async def generate_tender_docx(stage_outputs: dict, project_info: dict) -> bytes:
    """Generate a tender proposal Word document.

    Args:
        stage_outputs: Dict of all pipeline stage outputs (keyed by stage number).
        project_info: Project metadata (name, client_name, industry, etc.).

    Returns:
        .docx file as bytes.
    """
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "微软雅黑"
    font.size = Pt(10.5)

    # Heading styles
    for level in range(1, 4):
        h_style = doc.styles[f"Heading {level}"]
        h_style.font.name = "微软雅黑"
        h_style.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        if level == 1:
            h_style.font.size = Pt(18)
        elif level == 2:
            h_style.font.size = Pt(14)
        else:
            h_style.font.size = Pt(12)

    # ── Cover Page ──
    _add_cover_page(doc, project_info)

    # ── Table of Contents placeholder ──
    doc.add_page_break()
    doc.add_heading("目录", level=1)
    p = doc.add_paragraph()
    p.add_run("[请在 Word 中右键此处 → 更新域 以生成目录]").italic = True

    # ── Content Chapters ──
    requirements = stage_outputs.get(1, {})
    clarifications = stage_outputs.get(2, {})
    data_analysis = stage_outputs.get(3, {})
    knowledge = stage_outputs.get(4, {})
    solution = stage_outputs.get(5, {})
    automation = stage_outputs.get(6, {})
    benchmarks = stage_outputs.get(7, {})
    cost_model = stage_outputs.get(8, {})
    risks = stage_outputs.get(9, {})
    tender_content = stage_outputs.get(10, {})

    # If tender_writer produced chapter content, use it directly
    chapters = tender_content.get("document_structure", [])

    if chapters:
        for ch in chapters:
            doc.add_page_break()
            doc.add_heading(f"第{ch.get('chapter', '')}章 {ch.get('title', '')}", level=1)
            content = ch.get("content", "")
            _add_markdown_content(doc, content)
    else:
        # Fallback: build chapters from stage outputs
        _build_chapter_1(doc, project_info)
        _build_chapter_2(doc, requirements, clarifications)
        _build_chapter_3(doc, solution, data_analysis)
        _build_chapter_4(doc, automation)
        _build_chapter_5(doc, cost_model)
        _build_chapter_6(doc, risks)
        _build_chapter_7(doc, benchmarks)
        _build_chapter_8(doc, cost_model)

    # Add executive summary if available
    exec_summary = tender_content.get("executive_summary", "")
    if exec_summary:
        doc.add_page_break()
        doc.add_heading("附录：执行摘要", level=1)
        _add_markdown_content(doc, exec_summary)

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _add_cover_page(doc: Document, info: dict):
    """Create a professional cover page."""
    # Add spacing
    for _ in range(6):
        doc.add_paragraph()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(info.get("name", "物流解决方案投标文件"))
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("技术方案与商务报价")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph()

    # Info table
    info_items = [
        ("客户名称", info.get("client_name", "——")),
        ("所属行业", info.get("industry", "——")),
        ("编制日期", datetime.now().strftime("%Y年%m月%d日")),
        ("文档版本", "V1.0"),
        ("密级", "商业机密"),
    ]

    table = doc.add_table(rows=len(info_items), cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, (label, value) in enumerate(info_items):
        row = table.rows[i]
        cell_l = row.cells[0]
        cell_r = row.cells[1]
        cell_l.text = label
        cell_r.text = value
        _set_cell_shading(cell_l, "F1F5F9")
        for cell in (cell_l, cell_r):
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(11)


def _add_markdown_content(doc: Document, text: str):
    """Convert simple markdown-like text to Word paragraphs."""
    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped[0:3] in ("1. ", "2. ", "3. ", "4. ", "5. ", "6. ", "7. ", "8. ", "9. "):
            p = doc.add_paragraph(stripped[3:], style="List Number")
        elif stripped.startswith("**") and stripped.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(stripped[2:-2])
            run.bold = True
        else:
            doc.add_paragraph(stripped)


def _build_chapter_1(doc, info):
    doc.add_page_break()
    doc.add_heading("第一章 公司介绍与资质", level=1)
    doc.add_paragraph(
        f"本公司针对{info.get('industry', '物流')}行业客户提供一站式仓储物流解决方案，"
        "涵盖仓储运营、供应链管理、自动化集成等核心能力。"
    )


def _build_chapter_2(doc, requirements, clarifications):
    doc.add_page_break()
    doc.add_heading("第二章 项目理解与需求分析", level=1)

    reqs = requirements.get("requirements", [])
    if reqs:
        doc.add_heading("2.1 需求清单", level=2)
        table = doc.add_table(rows=1 + len(reqs[:15]), cols=4)
        table.style = "Table Grid"
        headers = ["编号", "需求描述", "优先级", "清晰度"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
            _set_cell_shading(table.rows[0].cells[i], "1F3864")
            for run in table.rows[0].cells[i].paragraphs[0].runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True

        for j, req in enumerate(reqs[:15]):
            row = table.rows[j + 1]
            row.cells[0].text = req.get("id", f"REQ-{j+1:03d}")
            row.cells[1].text = req.get("description", "")[:80]
            row.cells[2].text = req.get("priority", "P1")
            row.cells[3].text = req.get("clarity", "clear")

    clars = clarifications.get("clarifications_needed", [])
    if clars:
        doc.add_heading("2.2 需要澄清的事项", level=2)
        for c in clars[:10]:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"[{c.get('priority', 'P1')}] ")
            run.bold = True
            p.add_run(c.get("question", ""))


def _build_chapter_3(doc, solution, data_analysis):
    doc.add_page_break()
    doc.add_heading("第三章 物流解决方案设计", level=1)

    summary = solution.get("executive_summary", "")
    if summary:
        doc.add_paragraph(summary)

    # Warehouse design
    wh = solution.get("warehouse_design", {})
    if wh:
        doc.add_heading("3.1 仓储布局设计", level=2)
        doc.add_paragraph(f"总面积: {wh.get('total_area_sqm', '——')} ㎡")
        doc.add_paragraph(wh.get("flow_design", ""))

    # Operations
    ops = solution.get("operations_design", {})
    if ops:
        doc.add_heading("3.2 运营流程设计", level=2)
        for key, label in [("inbound", "入库"), ("picking", "拣选"), ("packing_shipping", "包装发运")]:
            section = ops.get(key, {})
            if section:
                doc.add_heading(f"3.2.x {label}流程", level=3)
                if isinstance(section, dict):
                    for k, v in section.items():
                        doc.add_paragraph(f"{k}: {v}")
                else:
                    doc.add_paragraph(str(section))


def _build_chapter_4(doc, automation):
    doc.add_page_break()
    doc.add_heading("第四章 自动化与技术方案", level=1)

    doc.add_paragraph(f"自动化等级: {automation.get('automation_level', '——')}")

    recs = automation.get("recommendations", [])
    if recs:
        doc.add_heading("4.1 推荐方案", level=2)
        for rec in recs:
            doc.add_heading(f"{rec.get('technology', '')} — {rec.get('application_area', '')}", level=3)
            doc.add_paragraph(f"适配评分: {rec.get('suitability_score', 0)}/10")
            doc.add_paragraph(f"投资估算: ¥{rec.get('estimated_cost_cny', 0):,.0f}")
            doc.add_paragraph(f"年节省: ¥{rec.get('annual_savings_cny', 0):,.0f}")
            doc.add_paragraph(f"ROI: {rec.get('roi_percent', 0)}%, 回本: {rec.get('payback_months', 0)}个月")
            doc.add_paragraph(rec.get("justification", ""))


def _build_chapter_5(doc, cost_model):
    doc.add_page_break()
    doc.add_heading("第五章 实施计划与里程碑", level=1)
    doc.add_paragraph("详细的项目实施计划将在合同签订后制定，主要包括：")
    for item in ["项目启动与团队组建", "场地准备与设备采购", "系统部署与集成", "试运行与调优", "正式运营与持续优化"]:
        doc.add_paragraph(item, style="List Number")


def _build_chapter_6(doc, risks):
    doc.add_page_break()
    doc.add_heading("第六章 风险管理与应急预案", level=1)

    risk_items = risks.get("risk_matrix", [])
    if risk_items:
        table = doc.add_table(rows=1 + len(risk_items[:10]), cols=4)
        table.style = "Table Grid"
        for i, h in enumerate(["风险描述", "可能性", "影响", "缓解措施"]):
            table.rows[0].cells[i].text = h
            _set_cell_shading(table.rows[0].cells[i], "1F3864")
            for run in table.rows[0].cells[i].paragraphs[0].runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True

        for j, r in enumerate(risk_items[:10]):
            row = table.rows[j + 1]
            row.cells[0].text = r.get("description", "")[:60]
            row.cells[1].text = r.get("likelihood", "")
            row.cells[2].text = r.get("impact", "")
            row.cells[3].text = r.get("mitigation", "")[:60]


def _build_chapter_7(doc, benchmarks):
    doc.add_page_break()
    doc.add_heading("第七章 成功案例参考", level=1)

    cases = benchmarks.get("matched_cases", [])
    for case in cases[:5]:
        doc.add_heading(case.get("case_name", "案例"), level=2)
        doc.add_paragraph(f"行业: {case.get('client_industry', '')}")
        doc.add_paragraph(f"相似度: {case.get('similarity_score', 0):.0%}")
        doc.add_paragraph(f"适用性: {case.get('applicable_to_current', '')}")


def _build_chapter_8(doc, cost_model):
    doc.add_page_break()
    doc.add_heading("第八章 报价方案", level=1)

    pricing = cost_model.get("pricing", {})
    if pricing:
        doc.add_heading("8.1 价格体系", level=2)
        items = [
            ("每单操作费", pricing.get("per_order"), "元/单"),
            ("每托仓储费", pricing.get("per_pallet"), "元/托"),
            ("场地月租", pricing.get("per_sqm_month"), "元/㎡/月"),
        ]
        for label, val, unit in items:
            if val:
                doc.add_paragraph(f"{label}: ¥{val:,.1f} {unit}")

    indicators = cost_model.get("financial_indicators", {})
    if indicators:
        doc.add_heading("8.2 投资回报分析", level=2)
        doc.add_paragraph(f"投资回报率 (ROI): {indicators.get('roi_percent', 0):.1f}%")
        doc.add_paragraph(f"内部收益率 (IRR): {indicators.get('irr_percent', 0):.1f}%")
        doc.add_paragraph(f"净现值 (NPV): ¥{indicators.get('npv_at_8pct', 0):,.0f}")
        doc.add_paragraph(f"回本周期: {indicators.get('payback_months', 0)} 个月")
