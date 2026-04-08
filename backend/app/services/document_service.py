"""Document processing service: extract text from PDF, Word, Excel, CSV files."""

import io
import structlog

logger = structlog.get_logger()


async def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """Extract text content from uploaded files.

    Supports: PDF, DOCX, DOC, TXT, XLSX, XLS, CSV, TSV
    """
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

    if ext == "pdf":
        return await _extract_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        return await _extract_docx(file_bytes)
    elif ext in ("xlsx", "xls"):
        return await _extract_excel(file_bytes, filename)
    elif ext == "csv":
        return await _extract_csv(file_bytes)
    elif ext == "tsv":
        return await _extract_csv(file_bytes, delimiter="\t")
    elif ext == "txt":
        return file_bytes.decode("utf-8", errors="replace")
    else:
        logger.warning("unsupported_file_type", ext=ext, filename=filename)
        return file_bytes.decode("utf-8", errors="replace")


async def _extract_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using pdfplumber."""
    import pdfplumber

    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for i, page in enumerate(pdf.pages):
            page_text = page.extract_text() or ""
            if page_text.strip():
                text_parts.append(f"--- Page {i + 1} ---\n{page_text}")

            # Also extract tables
            tables = page.extract_tables()
            for table in tables:
                if table:
                    rows = []
                    for row in table:
                        cells = [str(c or "").strip() for c in row]
                        rows.append(" | ".join(cells))
                    text_parts.append("TABLE:\n" + "\n".join(rows))

    return "\n\n".join(text_parts)


async def _extract_docx(file_bytes: bytes) -> str:
    """Extract text from Word documents."""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Extract tables too
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            paragraphs.append(" | ".join(cells))

    return "\n\n".join(paragraphs)


async def _extract_excel(file_bytes: bytes, filename: str = "") -> str:
    """Extract data from Excel files as structured text.

    Reads all sheets and converts to pipe-delimited text with headers.
    """
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True, read_only=True)
    parts = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c if c is not None else "").strip() for c in row]
            if any(cells):  # Skip fully empty rows
                rows.append(" | ".join(cells))

        if rows:
            parts.append(f"=== SHEET: {sheet_name} ({len(rows)} rows) ===")
            # Include header + first 500 data rows
            parts.extend(rows[:501])
            if len(rows) > 501:
                parts.append(f"... ({len(rows) - 501} more rows omitted)")

    wb.close()

    result = "\n".join(parts)
    logger.info("excel_extracted", filename=filename, sheets=len(wb.sheetnames), total_chars=len(result))
    return result


async def _extract_csv(file_bytes: bytes, delimiter: str = ",") -> str:
    """Extract data from CSV/TSV files."""
    import csv

    text = file_bytes.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text), delimiter=delimiter)

    rows = []
    for i, row in enumerate(reader):
        if i > 500:  # Limit to 500 rows
            rows.append(f"... (more rows omitted)")
            break
        rows.append(" | ".join(row))

    return f"=== CSV DATA ({len(rows)} rows) ===\n" + "\n".join(rows)
